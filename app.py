import streamlit as st
from pdf2docx import Converter
from docx import Document
import pytesseract
from PIL import Image
import os
import tempfile
from pdf2image import convert_from_path

# 1. Cấu hình trang
st.set_page_config(page_title="Học Liệu Số - Chuyển Đổi Pro", layout="wide")

# 2. Định nghĩa CSS để tạo giao diện Neon như ảnh mẫu
st.markdown("""
    <style>
    /* Màu nền tổng thể */
    .stApp {
        background-color: #00162b;
        color: #e0f2fe;
    }
    
    /* Container chính với viền phát sáng */
    .main-box {
        border: 2px solid #00d4ff;
        border-radius: 15px;
        padding: 20px;
        box-shadow: 0 0 15px #00d4ff;
        background-color: rgba(0, 40, 80, 0.5);
        margin-bottom: 20px;
    }
    
    /* Tiêu đề phong cách Cyberpunk */
    .header-title {
        text-align: center;
        color: #00d4ff;
        font-family: 'Segoe UI', sans-serif;
        text-transform: uppercase;
        letter-spacing: 2px;
        border: 2px solid #00d4ff;
        padding: 15px;
        border-radius: 10px;
        box-shadow: inset 0 0 10px #00d4ff;
        margin-bottom: 30px;
    }

    /* Các nút bấm tùy chỉnh */
    .stButton>button {
        background-color: transparent;
        color: #00d4ff;
        border: 1px solid #00d4ff;
        border-radius: 5px;
        transition: all 0.3s;
    }
    .stButton>button:hover {
        background-color: #00d4ff;
        color: #00162b;
        box-shadow: 0 0 20px #00d4ff;
    }

    /* Tabs tùy chỉnh */
    .stTabs [data-baseweb="tab-list"] {
        gap: 24px;
        background-color: transparent;
    }
    .stTabs [data-baseweb="tab"] {
        color: #00d4ff;
        font-weight: bold;
    }
    </style>
    """, unsafe_allow_html=True)

# --- GIAO DIỆN HEADER ---
st.markdown('<div class="header-title"><h1>CHUYỂN ĐỔI TÀI LIỆU SỐ</h1><p style="color:#00d4ff; font-size:14px;">Tiện ích được tạo bởi "Học liệu số"</p></div>', unsafe_allow_html=True)

# --- CHIA CỘT CHÍNH ---
col_left, col_right = st.columns([1, 1.5])

with col_left:
    st.markdown('<div class="main-box">', unsafe_allow_html=True)
    st.subheader("🛠️ BẢNG ĐIỀU KHIỂN")
    
    tab_type = st.radio("Chọn loại tài liệu:", ["📄 PDF", "🖼️ Ảnh (OCR)"])
    
    if tab_type == "📄 PDF":
        pdf_file = st.file_uploader("Tải lên PDF:", type="pdf")
        is_scan = st.toggle("Chế độ quét ảnh (Scan/OCR)", value=False)
    else:
        img_files = st.file_uploader("Tải lên một hoặc nhiều ảnh:", type=["jpg", "png", "jpeg"], accept_multiple_files=True)
    
    st.markdown('</div>', unsafe_allow_html=True)

with col_right:
    st.markdown('<div class="main-box" style="min-height: 400px;">', unsafe_allow_html=True)
    st.subheader("🔍 KẾT QUẢ XỬ LÝ")
    
    if tab_type == "📄 PDF" and pdf_file:
        st.write(f"Đang chờ lệnh cho: **{pdf_file.name}**")
        if st.button("BẮT ĐẦU CHUYỂN ĐỔI PDF"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(pdf_file.getvalue())
                tmp_path = tmp.name
            
            output_docx = tmp_path.replace(".pdf", ".docx")
            try:
                with st.spinner("Hệ thống đang xử lý..."):
                    if not is_scan:
                        cv = Converter(tmp_path)
                        cv.convert(output_docx)
                        cv.close()
                    else:
                        images = convert_from_path(tmp_path)
                        doc = Document()
                        for i, image in enumerate(images):
                            text = pytesseract.image_to_string(image, lang='vie')
                            doc.add_paragraph(text)
                            doc.add_page_break()
                        doc.save(output_docx)
                
                st.success("Hoàn tất chuyển đổi!")
                with open(output_docx, "rb") as f:
                    st.download_button("📥 TẢI FILE WORD VỀ MÁY", f, file_name=pdf_file.name.replace(".pdf", ".docx"))
            except Exception as e:
                st.error(f"Lỗi hệ thống: {str(e)}")
            finally:
                if os.path.exists(tmp_path): os.remove(tmp_path)
    
    elif tab_type == "🖼️ Ảnh (OCR)" and img_files:
        st.write(f"Số lượng ảnh đã chọn: {len(img_files)}")
        if st.button("QUÉT OCR VÀ KẾT XUẤT"):
            doc = Document()
            with st.spinner("Đang đọc chữ từ ảnh..."):
                for img_file in img_files:
                    text = pytesseract.image_to_string(Image.open(img_file), lang='vie')
                    doc.add_paragraph(text)
                    doc.add_page_break()
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, "rb") as f:
                        st.download_button("📥 TẢI FILE WORD ĐÃ QUÉT", f, file_name="Ket_qua_OCR.docx")
    else:
        st.info("👉 Hãy tải tệp tin ở cột bên trái để bắt đầu xử lý nội dung chi tiết.")
    
    st.markdown('</div>', unsafe_allow_html=True)

# --- FOOTER ---
st.markdown('<div style="text-align:center; border: 1px solid #00d4ff; padding:10px; border-radius:5px; margin-top:20px; font-size:12px;">Dữ liệu hỗ trợ chuyển đổi chuyên nghiệp - THCS Lê Thánh Tôn</div>', unsafe_allow_html=True)
